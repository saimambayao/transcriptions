#!/usr/bin/env python3
"""
Generate Cooperative Development Guidebook DOCX from markdown chapter files.
Usage: python3 generate-docx.py

Reads all markdown files in order and produces a professional Word document.
Requires: pip3 install python-docx markdown
"""

import os
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Install python-docx: pip3 install python-docx")
    exit(1)

# Configuration
SCRIPT_DIR = Path(__file__).parent
OUTPUT_FILE = SCRIPT_DIR / "Cooperative-Development-Guidebook-Bangsamoro.docx"

# Font settings
FONT_BODY = "Inter"
FONT_MONO = "Consolas"
COLOR_NAVY = RGBColor(0x1B, 0x36, 0x5D)
COLOR_SLATE = RGBColor(0x2C, 0x5F, 0x7C)
COLOR_BODY = RGBColor(0x2D, 0x2D, 0x2D)
COLOR_GOLD = RGBColor(0xC5, 0xA5, 0x4E)

# Source files in order
SOURCE_FILES = [
    "00-table-of-contents.md",
    "00b-about.md",
    "00c-author.md",
    "01-introduction.md",
    "02-chapter-01.md",
    "03-chapter-02.md",
    "04-chapter-03.md",
    "05-chapter-04.md",
    "06-chapter-05.md",
    "07-chapter-06.md",
    "08-chapter-07.md",
    "09-chapter-08.md",
    "10-chapter-09.md",
    "11-chapter-10.md",
    "12-chapter-11.md",
    "13-chapter-12.md",
    "14-chapter-13.md",
    "15-chapter-14.md",
    "16-glossary.md",
    "17-appendices.md",
]


def set_run_font(run, font_name=FONT_BODY, size=None, bold=False, italic=False, color=None):
    """Apply font settings to a run."""
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)

    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level, page_break=False):
    """Add a heading with custom formatting."""
    if page_break and level == 1:
        p = doc.add_paragraph()
        p.paragraph_format.page_break_before = True

    heading = doc.add_heading(text, level=level)

    for run in heading.runs:
        if level == 1:
            set_run_font(run, FONT_BODY, size=22, bold=True, color=COLOR_NAVY)
        elif level == 2:
            set_run_font(run, FONT_BODY, size=16, bold=True, color=COLOR_NAVY)
        elif level == 3:
            set_run_font(run, FONT_BODY, size=13, bold=True, color=COLOR_SLATE)
        elif level == 4:
            set_run_font(run, FONT_BODY, size=11.5, bold=True, italic=True, color=COLOR_SLATE)

    return heading


def add_body_paragraph(doc, text, indent=False, italic=False):
    """Add a body paragraph."""
    p = doc.add_paragraph()

    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, FONT_BODY, size=11, bold=True, color=COLOR_BODY)
        else:
            run = p.add_run(part)
            set_run_font(run, FONT_BODY, size=11, italic=italic, color=COLOR_BODY)

    if indent:
        p.paragraph_format.left_indent = Cm(1.5)

    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15

    return p


def add_blockquote(doc, text):
    """Add a blockquote (indented, italic)."""
    p = add_body_paragraph(doc, text, indent=True, italic=True)
    return p


def add_table_from_lines(doc, header_line, data_lines):
    """Add a formatted table."""
    def parse_row(line):
        cells = [c.strip() for c in line.strip('|').split('|')]
        return [c for c in cells if c]

    headers = parse_row(header_line)
    rows = [parse_row(line) for line in data_lines if not re.match(r'^[\s|:-]+$', line)]

    if not headers:
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        if i < len(table.columns):
            cell = table.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(header)
            set_run_font(run, FONT_BODY, size=9.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B365D"/>')
            cell._element.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < len(table.columns) and r_idx + 1 < len(table.rows):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = ""
                run = cell.paragraphs[0].add_run(cell_text)
                set_run_font(run, FONT_BODY, size=9.5, color=COLOR_BODY)
                if r_idx % 2 == 1:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F3F6"/>')
                    cell._element.get_or_add_tcPr().append(shading)

    doc.add_paragraph()


def process_markdown_file(doc, filepath, is_first_chapter=False):
    """Process a single markdown file and add its content to the document."""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_table = False
    table_header = None
    table_data = []
    in_blockquote = False
    blockquote_lines = []
    is_first_h1 = True

    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        if not stripped:
            if in_blockquote and blockquote_lines:
                text = ' '.join(blockquote_lines)
                add_blockquote(doc, text)
                blockquote_lines = []
                in_blockquote = False
            if in_table and table_header:
                add_table_from_lines(doc, table_header, table_data)
                table_header = None
                table_data = []
                in_table = False
            i += 1
            continue

        # Tables
        if '|' in stripped and not stripped.startswith('>'):
            if not in_table:
                in_table = True
                table_header = stripped
                table_data = []
            elif re.match(r'^[\s|:-]+$', stripped):
                pass
            else:
                table_data.append(stripped)
            i += 1
            continue
        elif in_table:
            add_table_from_lines(doc, table_header, table_data)
            table_header = None
            table_data = []
            in_table = False

        # Blockquotes
        if stripped.startswith('>'):
            in_blockquote = True
            text = stripped.lstrip('> ').strip()
            if text:
                blockquote_lines.append(text)
            i += 1
            continue
        elif in_blockquote:
            text = ' '.join(blockquote_lines)
            add_blockquote(doc, text)
            blockquote_lines = []
            in_blockquote = False

        # Headings
        if stripped.startswith('# ') and not stripped.startswith('## '):
            text = stripped[2:].strip()
            page_break = not is_first_h1 or not is_first_chapter
            add_heading(doc, text, 1, page_break=page_break)
            is_first_h1 = False
            i += 1
            continue
        elif stripped.startswith('## '):
            add_heading(doc, stripped[3:].strip(), 2)
            i += 1
            continue
        elif stripped.startswith('### '):
            add_heading(doc, stripped[4:].strip(), 3)
            i += 1
            continue
        elif stripped.startswith('#### '):
            add_heading(doc, stripped[5:].strip(), 4)
            i += 1
            continue

        # Horizontal rules
        if stripped in ('---', '***', '___'):
            i += 1
            continue

        # List items
        if re.match(r'^[-*]\s', stripped):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_run_font(run, FONT_BODY, size=11, bold=True, color=COLOR_BODY)
                else:
                    run = p.add_run(part)
                    set_run_font(run, FONT_BODY, size=11, color=COLOR_BODY)
            i += 1
            continue

        if re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_run_font(run, FONT_BODY, size=11, bold=True, color=COLOR_BODY)
                else:
                    run = p.add_run(part)
                    set_run_font(run, FONT_BODY, size=11, color=COLOR_BODY)
            i += 1
            continue

        # Regular paragraph
        text = stripped.replace('&emsp;', '    ')
        add_body_paragraph(doc, text)
        i += 1

    # Flush remaining
    if in_blockquote and blockquote_lines:
        add_blockquote(doc, ' '.join(blockquote_lines))
    if in_table and table_header:
        add_table_from_lines(doc, table_header, table_data)


def add_header_footer(doc):
    """Add running header and footer."""
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Cooperative Development Guidebook — BARMM")
        set_run_font(run, FONT_BODY, size=8, italic=True, color=COLOR_SLATE)

        # Footer with page number
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)
        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)
        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)
        set_run_font(run, FONT_BODY, size=9, color=COLOR_SLATE)


def create_title_page(doc):
    """Create a professional title page."""
    for _ in range(6):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("COOPERATIVE DEVELOPMENT")
    set_run_font(run, FONT_BODY, size=28, bold=True, color=COLOR_NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GUIDEBOOK")
    set_run_font(run, FONT_BODY, size=28, bold=True, color=COLOR_NAVY)

    doc.add_paragraph()

    # Gold separator
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2501" * 40)
    set_run_font(run, FONT_BODY, size=14, color=COLOR_GOLD)

    doc.add_paragraph()

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A Practical Guide for Community Organizers,\nCooperative Officers, and CSEA-BARMM Staff")
    set_run_font(run, FONT_BODY, size=14, italic=True, color=COLOR_SLATE)

    for _ in range(4):
        doc.add_paragraph()

    # Author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Saidamen R. Mambayao")
    set_run_font(run, FONT_BODY, size=12, bold=True, color=COLOR_NAVY)

    for _ in range(2):
        doc.add_paragraph()

    # Institution bottom
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Bangsamoro Government")
    set_run_font(run, FONT_BODY, size=12, bold=True, color=COLOR_NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Bangsamoro Government Center, Cotabato City")
    set_run_font(run, FONT_BODY, size=11, color=COLOR_BODY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Bangsamoro Autonomous Region in Muslim Mindanao")
    set_run_font(run, FONT_BODY, size=11, color=COLOR_BODY)

    # Page break after title page
    doc.add_page_break()


def main():
    doc = Document()

    # Set default margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Title page
    create_title_page(doc)

    # Process each source file
    for idx, filename in enumerate(SOURCE_FILES):
        filepath = SCRIPT_DIR / filename
        if not filepath.exists():
            print(f"Warning: {filename} not found, skipping")
            continue
        print(f"Processing: {filename}")
        process_markdown_file(doc, filepath, is_first_chapter=(idx == 0))

    # Add header/footer
    add_header_footer(doc)

    # Save
    doc.save(str(OUTPUT_FILE))
    print(f"\nSaved: {OUTPUT_FILE}")
    print(f"Size: {OUTPUT_FILE.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
