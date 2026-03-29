#!/usr/bin/env python3
"""
Generate a professional DOCX of the Manual of Operations for Bill Drafting
from the single Markdown source file.

Usage:
    python3 generate-mop-docx.py

Dependencies:
    pip3 install python-docx
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Mm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("[ERROR] python-docx not installed. Run: pip3 install python-docx")
    sys.exit(1)

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
BASE_DIR = Path(__file__).parent
SOURCE_MD = BASE_DIR / "Manual-of-Operations-Bill-Drafting.md"
OUTPUT_DOCX = BASE_DIR / "Manual-of-Operations-Bill-Drafting.docx"

# Colors
COLOR_NAVY = RGBColor(0x1B, 0x36, 0x5D)
COLOR_BODY = RGBColor(0x2D, 0x2D, 0x2D)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FONT_NAME = "Calibri"


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------
def set_run_font(run, size=11, bold=False, italic=False, color=COLOR_BODY, font=FONT_NAME):
    """Apply font settings to a run."""
    run.font.name = font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:ascii="{font}" w:hAnsi="{font}" w:cs="{font}"/>'
        )
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), font)
        rFonts.set(qn('w:hAnsi'), font)
        rFonts.set(qn('w:cs'), font)

    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_rich_text(paragraph, text, default_size=11, default_bold=False,
                  default_italic=False, default_color=COLOR_BODY):
    """Add text to a paragraph, handling **bold** and *italic* markdown markers."""
    # Handle bold+italic, bold, italic
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*[^*]+?\*)', text)
    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            set_run_font(run, size=default_size, bold=True, italic=True, color=default_color)
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=default_size, bold=True, italic=default_italic, color=default_color)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=default_size, bold=default_bold, italic=True, color=default_color)
        else:
            if part:
                run = paragraph.add_run(part)
                set_run_font(run, size=default_size, bold=default_bold,
                             italic=default_italic, color=default_color)


def add_heading(doc, text, level, page_break=False):
    """Add a formatted heading."""
    if page_break:
        # Add an empty paragraph with page break
        p = doc.add_paragraph()
        p.paragraph_format.page_break_before = True
        # Remove the empty paragraph spacing
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    heading = doc.add_heading(text, level=level)

    sizes = {1: 16, 2: 14, 3: 12, 4: 11}
    size = sizes.get(level, 11)

    for run in heading.runs:
        set_run_font(run, size=size, bold=True, color=COLOR_NAVY)

    return heading


def add_body_paragraph(doc, text, indent=False, italic=False):
    """Add a body paragraph with rich text support."""
    p = doc.add_paragraph()
    add_rich_text(p, text, default_italic=italic)
    if indent:
        p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_blockquote(doc, text):
    """Add a blockquote paragraph (indented, italic)."""
    return add_body_paragraph(doc, text, indent=True, italic=True)


def add_table(doc, header_line, data_lines):
    """Add a formatted table with Navy headers and alternating rows."""
    def parse_row(line):
        cells = [c.strip() for c in line.strip('|').split('|')]
        return cells

    headers = parse_row(header_line)
    rows = [parse_row(l) for l in data_lines if not re.match(r'^[\s|:-]+$', l)]

    if not headers:
        return

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row — Navy background, white text
    for i, header in enumerate(headers):
        if i < num_cols:
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            add_rich_text(p, header, default_size=9.5, default_bold=True, default_color=COLOR_WHITE)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B365D"/>')
            cell._element.get_or_add_tcPr().append(shading)

    # Data rows — alternating shading
    for r_idx, row_data in enumerate(rows):
        for c_idx in range(num_cols):
            cell_text = row_data[c_idx] if c_idx < len(row_data) else ""
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            add_rich_text(p, cell_text, default_size=9.5)
            # Alternating row shading
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F3F6"/>')
                cell._element.get_or_add_tcPr().append(shading)

    doc.add_paragraph()  # Space after table


def setup_document():
    """Create and configure the Document with margins, header, footer."""
    doc = Document()

    # Set margins: 25mm left/right, 20mm top, 25mm bottom
    for section in doc.sections:
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(25)
        section.right_margin = Mm(25)

    # Header — centered, 9pt
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run("Manual of Operations \u2014 Legislative Drafting")
    set_run_font(run, size=9, italic=True, color=RGBColor(0x8A, 0x8A, 0x8A))

    # Footer — left: institution, right: page number
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    # Create a tab stop for right-aligned page number
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run_left = fp.add_run("Bangsamoro Parliament")
    set_run_font(run_left, size=8, italic=True, color=RGBColor(0x99, 0x99, 0x99))

    # Add a right-aligned tab and page number using field code
    run_tab = fp.add_run("\t\t\t\t\t\t\t\t\t\t")
    set_run_font(run_tab, size=8)

    run_page = fp.add_run("Page ")
    set_run_font(run_page, size=8, color=RGBColor(0x66, 0x66, 0x66))

    # Insert PAGE field
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_fld = fp.add_run()
    run_fld._element.append(fld_char_begin)

    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run_instr = fp.add_run()
    run_instr._element.append(instr)

    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_end = fp.add_run()
    run_end._element.append(fld_char_end)

    return doc


def add_cover_page(doc):
    """Add a simple cover page."""
    # Add spacing before title
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

    # Institution
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BANGSAMORO PARLIAMENT")
    set_run_font(run, size=12, bold=False, color=COLOR_NAVY)
    run.font.all_caps = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Bangsamoro Autonomous Region in Muslim Mindanao")
    set_run_font(run, size=11, bold=False, color=COLOR_NAVY)

    # Spacing
    for _ in range(3):
        doc.add_paragraph()

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MANUAL OF OPERATIONS")
    set_run_font(run, size=28, bold=True, color=COLOR_NAVY)

    # Gold line (simulated with a thin paragraph)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 40)
    set_run_font(run, size=10, color=RGBColor(0xC5, 0xA5, 0x4E))

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Legislative Drafting for MP Offices")
    set_run_font(run, size=14, italic=True, color=RGBColor(0x4A, 0x4A, 0x4A))

    # Spacing
    for _ in range(6):
        doc.add_paragraph()

    # Author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Saidamen R. Mambayao")
    set_run_font(run, size=12, color=COLOR_BODY)

    # Page break after cover
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True


def process_markdown(doc, md_text):
    """Parse the markdown text and add content to the document."""
    lines = md_text.split('\n')
    i = 0
    is_first_h1 = True

    # State tracking
    in_table = False
    table_header = None
    table_data = []
    in_blockquote = False
    blockquote_lines = []
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        # Code blocks
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(code_text)
                set_run_font(run, size=9, font="Consolas", color=COLOR_BODY)
                # Add shading
                pPr = p._element.get_or_add_pPr()
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F3F6"/>')
                pPr.append(shading)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Empty lines
        if not stripped:
            if in_blockquote and blockquote_lines:
                text = ' '.join(blockquote_lines)
                add_blockquote(doc, text)
                blockquote_lines = []
                in_blockquote = False
            if in_table and table_header:
                add_table(doc, table_header, table_data)
                table_header = None
                table_data = []
                in_table = False
            i += 1
            continue

        # Tables
        if '|' in stripped and not stripped.startswith('>'):
            if not in_table:
                in_table = True
                table_header = stripped
                table_data = []
            elif re.match(r'^[\s|:-]+$', stripped):
                pass  # Separator line
            else:
                table_data.append(stripped)
            i += 1
            continue
        elif in_table:
            add_table(doc, table_header, table_data)
            table_header = None
            table_data = []
            in_table = False

        # Blockquotes
        if stripped.startswith('>'):
            in_blockquote = True
            text = stripped.lstrip('> ').strip()
            if text:
                blockquote_lines.append(text)
            i += 1
            continue
        elif in_blockquote:
            text = ' '.join(blockquote_lines)
            add_blockquote(doc, text)
            blockquote_lines = []
            in_blockquote = False

        # Headings
        if stripped.startswith('# ') and not stripped.startswith('## '):
            text = stripped[2:].strip()
            page_break = not is_first_h1
            add_heading(doc, text, 1, page_break=page_break)
            is_first_h1 = False
            i += 1
            continue
        elif stripped.startswith('## '):
            add_heading(doc, stripped[3:].strip(), 2)
            i += 1
            continue
        elif stripped.startswith('### '):
            add_heading(doc, stripped[4:].strip(), 3)
            i += 1
            continue
        elif stripped.startswith('#### '):
            add_heading(doc, stripped[5:].strip(), 4)
            i += 1
            continue

        # Horizontal rules
        if stripped in ('---', '***', '___'):
            i += 1
            continue

        # Bullet list items (with sub-items)
        bullet_match = re.match(r'^(\s*)[-*]\s(.+)', line)
        if bullet_match:
            indent_level = len(bullet_match.group(1))
            text = bullet_match.group(2).strip()
            p = doc.add_paragraph(style='List Bullet')
            add_rich_text(p, text)
            if indent_level >= 2:
                p.paragraph_format.left_indent = Cm(2.5)
            i += 1
            continue

        # Numbered list items
        num_match = re.match(r'^(\s*)\d+\.\s(.+)', line)
        if num_match:
            text = num_match.group(2).strip()
            p = doc.add_paragraph(style='List Number')
            add_rich_text(p, text)
            i += 1
            continue

        # Regular paragraph
        text = stripped.replace('&emsp;', '    ')
        add_body_paragraph(doc, text)
        i += 1

    # Flush remaining
    if in_blockquote and blockquote_lines:
        text = ' '.join(blockquote_lines)
        add_blockquote(doc, text)
    if in_table and table_header:
        add_table(doc, table_header, table_data)


def main():
    print("=" * 60)
    print("  Manual of Operations — DOCX Generator")
    print("=" * 60)

    # 1. Read source markdown
    if not SOURCE_MD.exists():
        print(f"[ERROR] Source file not found: {SOURCE_MD}")
        sys.exit(1)

    md_text = SOURCE_MD.read_text(encoding="utf-8")
    print(f"  [OK] Read source: {SOURCE_MD.name} ({len(md_text):,} chars)")

    # 2. Create document
    doc = setup_document()
    print(f"  [OK] Document created with margins and header/footer")

    # 3. Add cover page
    add_cover_page(doc)
    print(f"  [OK] Cover page added")

    # 4. Process markdown content
    process_markdown(doc, md_text)
    print(f"  [OK] Markdown content processed")

    # 5. Save
    doc.save(str(OUTPUT_DOCX))
    print(f"  [OK] DOCX saved: {OUTPUT_DOCX}")

    # 6. Report size
    docx_size = OUTPUT_DOCX.stat().st_size
    if docx_size > 1_048_576:
        print(f"  [OK] DOCX size: {docx_size / 1_048_576:.1f} MB")
    else:
        print(f"  [OK] DOCX size: {docx_size / 1024:.0f} KB")

    print("=" * 60)
    print("  Done!")
    print("=" * 60)


if __name__ == "__main__":
    main()
