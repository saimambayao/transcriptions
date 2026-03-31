#!/usr/bin/env python3
"""
Convert Markdown files to Word (.docx) format for BARMM parliamentary documents.

Formatting specifications:
- Font: Bookman Old Style, 12pt
- Single spacing
- Standard margins (1 inch)
- Page numbers in footer
- WHEREAS clauses: Indented
- Footnote citations: Superscript maintained
- ALL CAPS titles: Preserved
- Bold text: Preserved
"""

import re
import sys
from pathlib import Path
from typing import List, Tuple

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx library not installed.")
    print("Install with: pip install python-docx")
    sys.exit(1)


class MarkdownToDocxConverter:
    """Convert Markdown to Word with BARMM parliamentary formatting."""

    # Superscript number mapping
    SUPERSCRIPT_MAP = {
        '¹': '1', '²': '2', '³': '3', '⁴': '4', '⁵': '5',
        '⁶': '6', '⁷': '7', '⁸': '8', '⁹': '9', '⁰': '0',
        '¹⁰': '10', '¹¹': '11', '¹²': '12', '¹³': '13', '¹⁴': '14',
        '¹⁵': '15', '¹⁶': '16', '¹⁷': '17', '¹⁸': '18', '¹⁹': '19',
        '²⁰': '20', '²¹': '21', '²²': '22', '²³': '23', '²⁴': '24',
        '²⁵': '25', '²⁶': '26', '²⁷': '27', '²⁸': '28', '²⁹': '29',
        '³⁰': '30', '³¹': '31', '³²': '32', '³³': '33', '³⁴': '34',
        '³⁵': '35', '³⁶': '36', '³⁷': '37', '³⁸': '38', '³⁹': '39',
        '⁴⁰': '40', '⁴¹': '41', '⁴²': '42', '⁴³': '43', '⁴⁴': '44',
        '⁴⁵': '45', '⁴⁶': '46', '⁴⁷': '47', '⁴⁸': '48', '⁴⁹': '49',
        '⁵⁰': '50'
    }

    def __init__(self):
        self.doc = Document()
        self.setup_document_formatting()

    def setup_document_formatting(self):
        """Set up document-wide formatting: margins, font, spacing."""
        # Set margins to 1 inch (standard)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

            # Add page numbers to footer
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.add_page_number(footer_para)

        # Set default font for Normal style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Bookman Old Style'
        font.size = Pt(12)

        # Set paragraph spacing
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing = 1.0  # Single spacing

    def add_page_number(self, paragraph):
        """Add page number field to paragraph."""
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def parse_inline_formatting(self, text: str) -> List[Tuple[str, dict]]:
        """
        Parse inline formatting (bold, superscripts) from text.
        Returns list of (text, formatting) tuples.
        """
        segments = []
        current_pos = 0

        # Pattern to find bold text (**text**) and superscript numbers
        bold_pattern = r'\*\*([^*]+)\*\*'
        superscript_pattern = r'([¹²³⁴⁵⁶⁷⁸⁹⁰]+)'

        # Combine patterns
        combined_pattern = f'({bold_pattern}|{superscript_pattern})'

        for match in re.finditer(combined_pattern, text):
            # Add text before match
            if match.start() > current_pos:
                plain_text = text[current_pos:match.start()]
                if plain_text:
                    segments.append((plain_text, {}))

            matched_text = match.group(0)

            # Check if it's bold
            if matched_text.startswith('**') and matched_text.endswith('**'):
                bold_text = matched_text[2:-2]
                segments.append((bold_text, {'bold': True}))
            # Check if it's superscript
            elif any(char in self.SUPERSCRIPT_MAP for char in matched_text):
                segments.append((matched_text, {'superscript': True}))
            else:
                segments.append((matched_text, {}))

            current_pos = match.end()

        # Add remaining text
        if current_pos < len(text):
            remaining_text = text[current_pos:]
            if remaining_text:
                segments.append((remaining_text, {}))

        return segments if segments else [(text, {})]

    def add_formatted_text(self, paragraph, text: str, base_formatting: dict = None):
        """Add text to paragraph with inline formatting preserved."""
        base_formatting = base_formatting or {}
        segments = self.parse_inline_formatting(text)

        for segment_text, segment_format in segments:
            run = paragraph.add_run(segment_text)

            # Apply base formatting
            run.font.name = 'Bookman Old Style'
            run.font.size = Pt(12)

            # Apply segment-specific formatting
            if segment_format.get('bold') or base_formatting.get('bold'):
                run.bold = True

            if segment_format.get('superscript'):
                run.font.superscript = True

            if base_formatting.get('all_caps'):
                run.font.all_caps = True

    def is_whereas_or_resolved(self, line: str) -> bool:
        """Check if line starts with WHEREAS or RESOLVED."""
        stripped = line.strip()
        return (stripped.startswith('WHEREAS,') or
                stripped.startswith('**WHEREAS,') or
                stripped.startswith('NOW, THEREFORE') or
                stripped.startswith('RESOLVED') or
                stripped.startswith('**RESOLVED'))

    def is_all_caps_title(self, line: str) -> bool:
        """Check if line is an ALL CAPS title."""
        stripped = line.strip()
        if not stripped:
            return False
        # Remove markdown bold markers
        cleaned = stripped.replace('**', '')
        # Check if mostly uppercase (allowing for some punctuation)
        words = re.findall(r'[A-Za-z]+', cleaned)
        if not words:
            return False
        upper_count = sum(1 for word in words if word.isupper())
        return upper_count / len(words) >= 0.7

    def process_markdown_line(self, line: str):
        """Process a single line of markdown and add to document."""
        stripped = line.strip()

        # Skip empty lines (but add paragraph break)
        if not stripped:
            self.doc.add_paragraph()
            return

        # Skip markdown code blocks
        if stripped.startswith('```'):
            return

        # Handle headers (## Title)
        if stripped.startswith('#'):
            level = len(re.match(r'^#+', stripped).group(0))
            text = stripped.lstrip('#').strip()

            if level == 1:
                para = self.doc.add_heading(text, level=1)
            elif level == 2:
                para = self.doc.add_heading(text, level=2)
            else:
                para = self.doc.add_heading(text, level=3)

            # Apply font formatting to headings
            for run in para.runs:
                run.font.name = 'Bookman Old Style'
                run.font.size = Pt(12)
                run.bold = True
            return

        # Handle WHEREAS and RESOLVED clauses (indented)
        if self.is_whereas_or_resolved(stripped):
            para = self.doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(0)
            self.add_formatted_text(para, stripped, {'bold': True})
            return

        # Handle ALL CAPS titles (centered)
        if self.is_all_caps_title(stripped):
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.add_formatted_text(para, stripped, {'bold': True, 'all_caps': True})
            return

        # Handle numbered lists
        if re.match(r'^\d+\.', stripped):
            para = self.doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\.\s*', '', stripped)
            self.add_formatted_text(para, text)
            return

        # Handle bullet lists
        if stripped.startswith('- ') or stripped.startswith('* '):
            para = self.doc.add_paragraph(style='List Bullet')
            text = stripped[2:]
            self.add_formatted_text(para, text)
            return

        # Handle regular paragraphs
        para = self.doc.add_paragraph()
        self.add_formatted_text(para, stripped)

    def convert(self, markdown_file: Path, output_file: Path):
        """Convert markdown file to Word document."""
        # Read markdown file
        with open(markdown_file, 'r', encoding='utf-8') as f:
            content = f.read()

        # Process line by line
        lines = content.split('\n')
        for line in lines:
            self.process_markdown_line(line)

        # Save Word document
        self.doc.save(output_file)
        print(f"✅ Converted: {markdown_file.name} → {output_file.name}")


def main():
    """Main conversion function."""
    if len(sys.argv) < 2:
        print("Usage: python convert_to_docx.py <markdown_file.md>")
        print("   or: python convert_to_docx.py <markdown_file.md> <output_file.docx>")
        sys.exit(1)

    input_file = Path(sys.argv[1])

    if not input_file.exists():
        print(f"Error: File not found: {input_file}")
        sys.exit(1)

    if not input_file.suffix == '.md':
        print(f"Error: Input file must be a .md (markdown) file")
        sys.exit(1)

    # Determine output file
    if len(sys.argv) >= 3:
        output_file = Path(sys.argv[2])
    else:
        output_file = input_file.with_suffix('.docx')

    # Convert
    converter = MarkdownToDocxConverter()
    converter.convert(input_file, output_file)
    print(f"✅ Word document saved to: {output_file}")


if __name__ == "__main__":
    main()
